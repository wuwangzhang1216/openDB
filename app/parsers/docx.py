"""DOCX parser using python-docx.

Splits on page breaks, extracts headings, tables, and metadata.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.parsers.base import Page, ParseResult
from app.parsers.registry import register

CHUNK_SIZE = 3000


class DocxParser:
    def parse(self, file_path: Path) -> ParseResult:
        doc = Document(str(file_path))

        # Collect all content elements in order (paragraphs + tables)
        elements = self._collect_elements(doc)

        # Split into pages by page breaks
        page_groups = self._split_by_page_breaks(elements)

        # If no page breaks found, chunk by size
        if len(page_groups) == 1 and self._total_len(page_groups[0]) > CHUNK_SIZE:
            page_groups = self._chunk_elements(page_groups[0])

        pages: list[Page] = []
        for i, group in enumerate(page_groups, start=1):
            text = "\n\n".join(group)
            title = self._extract_title_from_group(doc, i - 1, group)
            pages.append(
                Page(page_number=i, section_title=title, text=text.strip())
            )

        metadata = self._extract_metadata(doc)
        return ParseResult(pages=pages, extracted_metadata=metadata)

    def _collect_elements(self, doc: Document) -> list[tuple[str, str]]:
        """Collect (type, text) tuples for paragraphs and tables."""
        elements: list[tuple[str, str]] = []

        for element in doc.element.body:
            if element.tag.endswith("}p"):  # paragraph
                # Find corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element is element:
                        text = para.text.strip()
                        style = para.style.name if para.style else ""
                        has_break = self._has_page_break(para)
                        elem_type = "break" if has_break else "para"
                        if style.startswith("Heading"):
                            elem_type = f"heading:{style}"
                        if text or has_break:
                            elements.append((elem_type, text))
                        break
            elif element.tag.endswith("}tbl"):  # table
                for table in doc.tables:
                    if table._element is element:
                        table_text = self._format_table(table)
                        if table_text:
                            elements.append(("table", table_text))
                        break

        return elements

    def _has_page_break(self, para) -> bool:
        """Check if paragraph contains a page break."""
        for run in para.runs:
            if run._element.findall(qn("w:br")):
                for br in run._element.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        return True
        return False

    def _split_by_page_breaks(
        self, elements: list[tuple[str, str]]
    ) -> list[list[str]]:
        """Split elements into groups by page breaks."""
        groups: list[list[str]] = [[]]

        for elem_type, text in elements:
            if elem_type == "break" and groups[-1]:
                groups.append([])
            if text:
                groups[-1].append(text)

        # Filter empty groups
        return [g for g in groups if g]

    def _chunk_elements(self, texts: list[str]) -> list[list[str]]:
        """Chunk text list into ~CHUNK_SIZE groups."""
        groups: list[list[str]] = [[]]
        current_len = 0

        for text in texts:
            if current_len + len(text) > CHUNK_SIZE and groups[-1]:
                groups.append([])
                current_len = 0
            groups[-1].append(text)
            current_len += len(text) + 2

        return [g for g in groups if g]

    def _total_len(self, texts: list[str]) -> int:
        return sum(len(t) for t in texts)

    def _extract_title_from_group(
        self, doc: Document, group_idx: int, group: list[str]
    ) -> str | None:
        """Extract heading from the first few elements of a group."""
        for para in doc.paragraphs:
            if para.style and para.style.name.startswith("Heading"):
                if para.text.strip() in group[:3]:
                    return para.text.strip()
        return None

    def _format_table(self, table) -> str:
        """Format a DOCX table as pipe-delimited text."""
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if len(rows) > 1:
            header_cells = table.rows[0].cells
            separator = "|" + "|".join(
                "-" * (len(c.text.strip()) + 2) for c in header_cells
            ) + "|"
            rows.insert(1, separator)

        return "\n".join(rows)

    def _extract_metadata(self, doc: Document) -> dict:
        """Extract DOCX core properties."""
        props = doc.core_properties
        result = {}

        if props.author:
            result["author"] = props.author
        if props.title:
            result["title"] = props.title
        if hasattr(props, "company") and getattr(props, "company", None):
            result["company"] = props.company
        if props.created:
            result["created"] = props.created.isoformat()
        if props.modified:
            result["modified"] = props.modified.isoformat()
        if props.revision:
            result["revision"] = props.revision

        return result


register(
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    DocxParser(),
)
# Legacy .doc format — same parser may not fully work, but register anyway
register("application/msword", DocxParser())
